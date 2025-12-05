import os
import sys
from flask import Flask, render_template, send_from_directory, jsonify

# 檢查 docx 依賴
try:
    from docx import Document
except ImportError:
    print("❌ 錯誤：找不到必要套件 'python-docx'。")
    print("請在終端機執行指令: pip install -r requirements.txt")
    sys.exit(1)

# 🚨 部署修復 1: 明確指定 static_folder 確保在 Gunicorn 環境下靜態資源路徑正確 🚨
app = Flask(__name__, static_folder='static') 

# ================= 設定區域 =================
BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')

# 🚨 固定順序設置 🚨
LOCATION_ORDER = ["西門", "板橋", "中壢", "桃園", "聯絡我們"]

ALLOWED_EXTENSIONS = {
    'image': ['.jpg', '.jpeg', '.png', '.gif', '.webp'],
    'video': ['.mp4', '.mov', '.webm'],
    'text': ['.docx'] 
}

# ================= 輔助功能 (DOCX 處理) =================

def extract_preview(path):
    """提取 DOCX 文件的前三行文字作為預覽"""
    try:
        doc = Document(path)
        txt = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(txt[:3]) if txt else '尚無文字簡介'
    except: return '預覽讀取失敗'

def read_full_docx(path):
    """讀取 DOCX 文件的完整內容，用於內容詳情頁"""
    try:
        doc = Document(path)
        full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        # 確保預覽文字長度不會超過 200 字
        preview_text = full_text[:200] 
        return {
            'preview': preview_text,
            'full': full_text,
            'has_doc': True
        }
    except: 
        return {'preview': '內容讀取失敗', 'full': '內容讀取失敗', 'has_doc': False}

def read_full_docx_text(path):
    """讀取 DOCX 文件的純文本內容，用於聯絡資訊彈窗"""
    try:
        doc = Document(path)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    except: return '內容讀取失敗'


# ================= 路由邏輯 (API Endpoints) =================

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/locations')
def get_locations():
    if not os.path.exists(BASE_DIR):
        return jsonify([])
    
    # 獲取實際存在的目錄並按照 LOCATION_ORDER 排序
    existing_dirs = {d for d in os.listdir(BASE_DIR) if os.path.isdir(os.path.join(BASE_DIR, d))}
    sorted_locations = [loc for loc in LOCATION_ORDER if loc in existing_dirs]
    
    return jsonify(sorted_locations)

@app.route('/api/people/<location>')
def get_people(location):
    location_path = os.path.join(BASE_DIR, location)
    if not os.path.exists(location_path):
        return jsonify([])
    
    people_list = []
    
    person_dirs = []
    for entry in os.scandir(location_path):
        if entry.is_dir():
            dir_name = entry.name
            # 排除 contact 資料夾和隱藏資料夾
            if dir_name.lower() == 'contact' or dir_name.startswith('.'):
                continue
            else:
                person_dirs.append(dir_name)

    for person in person_dirs:
        person_path = os.path.join(location_path, person)
        p_info = {'name': person, 'thumbnail': None, 'preview': '尚無文字簡介'}
        
        try:
            if not os.listdir(person_path): continue 

            docx_file = next((f for f in os.listdir(person_path) if f.endswith('.docx')), None)
            if docx_file:
                p_info['preview'] = extract_preview(os.path.join(person_path, docx_file))
            
            thumbnail_file = next((f for f in os.listdir(person_path) 
                                   if os.path.splitext(f)[1].lower() in ALLOWED_EXTENSIONS['image']), None)
            if thumbnail_file:
                 p_info['thumbnail'] = f"/files/{location}/{person}/{thumbnail_file}"
        
        except Exception:
            continue
            
        people_list.append(p_info)
        
    return jsonify(people_list)


@app.route('/api/content/<location>/<person>')
def get_content(location, person):
    person_path = os.path.join(BASE_DIR, location, person)
    if not os.path.exists(person_path):
        return jsonify({'text': {'preview': '無內容', 'full': '無內容', 'has_doc': False}, 'images': [], 'videos': []})

    content = {'images': [], 'videos': [], 'text': {'preview': '無內容', 'full': '無內容', 'has_doc': False}}
    try:
        for file in os.listdir(person_path):
            ext = os.path.splitext(file)[1].lower()
            url = f"/files/{location}/{person}/{file}"
            
            if ext in ALLOWED_EXTENSIONS['image']:
                content['images'].append({'name': file, 'url': url})
            elif ext in ALLOWED_EXTENSIONS['video']:
                content['videos'].append({'name': file, 'url': url})
            elif ext in ALLOWED_EXTENSIONS['text']:
                text_data = read_full_docx(os.path.join(person_path, file))
                if text_data: content['text'] = text_data
    except Exception: 
        pass
    return jsonify(content)

# 聯絡資訊 API 路由
@app.route('/api/contact/<location>/<person>')
def get_contact_info(location, person):
    if person == '_location_':
        contact_path = os.path.join(BASE_DIR, location, 'contact')
        file_url_base = f"/files/{location}/_location_/contact"
    else:
        contact_path = os.path.join(BASE_DIR, location, person, 'contact')
        file_url_base = f"/files/{location}/{person}/contact"

    if not os.path.exists(contact_path):
        return jsonify({'images': [], 'text': []})

    contact_data = {'images': [], 'text': []}
    try:
        for file in os.listdir(contact_path):
            ext = os.path.splitext(file)[1].lower()
            url = f"{file_url_base}/{file}"
            
            if ext in ALLOWED_EXTENSIONS['image']:
                name = os.path.splitext(file)[0].upper()
                contact_data['images'].append({'name': name, 'url': url})
            elif ext in ALLOWED_EXTENSIONS['text']:
                text_content = read_full_docx_text(os.path.join(contact_path, file))
                if text_content:
                    name = os.path.splitext(file)[0].upper()
                    contact_data['text'].append({'name': name, 'content': text_content})
    except Exception as e:
        print(f"Error reading contact info: {e}")
        pass
    return jsonify(contact_data)

# 檔案服務路由 (給縮圖、圖片、影片使用)
@app.route('/files/<location>/<person>/<filename>')
def serve_file(location, person, filename):
    return send_from_directory(os.path.join(BASE_DIR, location, person), filename)

# 檔案服務路由 (給聯絡資訊檔案使用)
@app.route('/files/<location>/<person_or_location_tag>/contact/<filename>')
def serve_contact_file(location, person_or_location_tag, filename):
    if person_or_location_tag == '_location_':
        base = os.path.join(BASE_DIR, location, 'contact')
    else:
        base = os.path.join(BASE_DIR, location, person_or_location_tag, 'contact')
        
    return send_from_directory(base, filename)


# 本地啟動區塊 (部署環境下會使用 Gunicorn 運行)
if __name__ == '__main__':
    if not os.path.exists(BASE_DIR):
        os.makedirs(BASE_DIR)
        
    print(f"Flask 應用程式啟動中，數據目錄: {BASE_DIR}")
    
    # 🚨 部署修復 2: 使用動態端口 🚨
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)