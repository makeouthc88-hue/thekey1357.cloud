import os
from flask import Flask, render_template, send_from_directory, jsonify

# 初始化 Flask
app = Flask(__name__, template_folder='templates')

# ================= 配置區域 (雲端修正版) =================

# 取得目前這支程式 (flask_app.py) 所在的資料夾路徑
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))

# 設定資料根目錄
# 請務必將您的「情熱天堂班表」資料夾改名為 "data" 並放在跟 flask_app.py 同一層
BASE_DIR = os.path.join(CURRENT_DIR, 'data')

# =======================================================

ALLOWED_EXTENSIONS = {
    'image': ['.jpg', '.jpeg', '.png', '.gif', '.webp'],
    'video': ['.mp4', '.mov', '.webm'],
    'text': ['.docx'] 
}

# 嘗試匯入 docx，如果失敗則略過 (避免雲端部署直接報錯 Crash)
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ 警告: 未安裝 python-docx，將無法預覽 Word 檔。請在 Console 執行: pip install python-docx")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/locations')
def get_locations():
    """取得所有地點"""
    if not os.path.exists(BASE_DIR):
        # 如果找不到資料夾，回傳空陣列並印出路徑方便除錯
        print(f"❌ 錯誤：找不到資料路徑 {BASE_DIR}")
        return jsonify([])
    
    # 掃描資料夾
    dirs = [d for d in os.listdir(BASE_DIR) if os.path.isdir(os.path.join(BASE_DIR, d))]
    return jsonify(dirs)

@app.route('/api/people/<location>')
def get_people(location):
    """取得特定地點下的人員"""
    loc_path = os.path.join(BASE_DIR, location)
    if not os.path.exists(loc_path):
        return jsonify([])
    
    people_dirs = [d for d in os.listdir(loc_path) if os.path.isdir(os.path.join(loc_path, d))]
    people_data = []

    for p in people_dirs:
        p_path = os.path.join(loc_path, p)
        p_info = {
            'name': p,
            'thumbnail': None,
            'preview': '尚無簡介'
        }
        
        try:
            files = sorted(os.listdir(p_path))
            
            # 找封面圖
            for f in files:
                ext = os.path.splitext(f)[1].lower()
                if ext in ALLOWED_EXTENSIONS['image']:
                    # 注意：這裡的路徑要對應下面的 /files/ 路由
                    p_info['thumbnail'] = f"/files/{location}/{p}/{f}"
                    break
            
            # 找文字預覽 (如果有安裝 docx)
            if HAS_DOCX:
                for f in files:
                    if f.lower().endswith('.docx'):
                        try:
                            doc = Document(os.path.join(p_path, f))
                            lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                            if lines:
                                p_info['preview'] = '\n'.join(lines[:3])
                            break
                        except Exception:
                            pass
        except Exception as e:
            print(f"Error scanning {p}: {e}")

        people_data.append(p_info)

    return jsonify(people_data)

@app.route('/api/content/<location>/<person>')
def get_content(location, person):
    """讀取詳細內容"""
    person_path = os.path.join(BASE_DIR, location, person)
    if not os.path.exists(person_path):
        return jsonify({'error': 'Not found'}), 404

    content = {
        'images': [],
        'videos': [],
        'text': {'preview': '', 'full': '', 'has_doc': False}
    }

    try:
        for file in os.listdir(person_path):
            ext = os.path.splitext(file)[1].lower()
            file_url = f"/files/{location}/{person}/{file}"
            
            if ext in ALLOWED_EXTENSIONS['image']:
                content['images'].append({'name': file, 'url': file_url})
            elif ext in ALLOWED_EXTENSIONS['video']:
                content['videos'].append({'name': file, 'url': file_url})
            elif ext in ALLOWED_EXTENSIONS['text'] and HAS_DOCX:
                try:
                    doc = Document(os.path.join(person_path, file))
                    full_text = [p.text for p in doc.paragraphs if p.text.strip()]
                    content['text']['full'] = '\n'.join(full_text)
                    content['text']['preview'] = '\n'.join(full_text[:5])
                    content['text']['has_doc'] = True
                except Exception:
                    pass
    except Exception as e:
        print(f"Error reading content: {e}")

    return jsonify(content)

@app.route('/files/<location>/<person>/<filename>')
def serve_file(location, person, filename):
    """檔案伺服器"""
    directory = os.path.join(BASE_DIR, location, person)
    return send_from_directory(directory, filename)

if __name__ == '__main__':
    app.run(debug=True)